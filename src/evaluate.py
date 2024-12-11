# 导入所需的库和模块
from src import api_doc
from src import prompt_factor
import mosestokenizer 
from src import word_reader, utils
from docx import Document
from src import word_check
from sacremoses import MosesTokenizer
import os
from tqdm import tqdm
from src import utils, word_reader

def calc_token_cost(path):
    """计算文本的token数量"""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
        return len(text.split())
    except Exception as e:
        print(f"Error calculating token cost: {e}")
        return 0

def calc_acc(label_path, pred_path, instruction, additional_restrictions=[]):
    """计算准确率
    
    Args:
        label_path: 标准文档路径
        pred_path: 预测文档路径
        instruction: 指令文本
        additional_restrictions: 额外的限制条件
    
    Returns:
        str_correct: 文本内容是否正确
        pos_total: 位置要求总数
        pos_correct: 正确的位置要求数
    """
    pos_total, pos_correct, str_correct = 0, 0, 0
    
    # 处理位置相关的限制条件
    splitted = instruction.split('##')
    instruction, restrictions = splitted[0], splitted[1:]
    restrictions += additional_restrictions
    
    if len(restrictions) > 0:
        pos_total = 1
        pos_correct = 1
        
    try:
        # 加载文档
        pred_doc = Document(pred_path)
        label_doc = Document(label_path)
        
        # 检查位置要求
        for res in restrictions:
            para_id, element_type, position = [x.strip() for x in res.split(",")]
            try:
                pred_para = pred_doc.paragraphs[int(para_id)]
                pos_correct *= word_reader.check_element_position(pred_para, element_type, position)
            except Exception as e:
                print(f"Error checking position: {res}")
                print(f"Instruction: {instruction}")
                pos_correct = 0
        
        # 检查文本内容
        label_content = word_reader.get_document_content(label_doc)
        pred_content = word_reader.get_document_content(pred_doc)
        
        if label_content == pred_content:
            if len(restrictions) > 0:
                if pos_correct == 1:
                    str_correct = 1
            else:
                str_correct = 1
                
        print(f'String correct: {str_correct}')
        
    except Exception as e:
        print(f'Document Loading Failed: {e}')
        str_correct, pos_correct = 0, 0
    
    return str_correct, pos_total, pos_correct

def eval(args):
    """评估模型性能
    
    Args:
        args: 参数配置对象
    """
    set_name = 'Create_new_docs' if args.dataset == 'short' else 'Edit_Word_template'
    token_costs, api_costs = [], []
    string_total = string_correct = position_total = position_correct = 0
    
    if args.tf:  # 按turn评估
        turn_nums = []
        for sess_id, session_path in enumerate(utils.sorted_list(f'{args.user_path}/Word_test_output/{set_name}')):
            if not session_path.startswith(args.exp_name):
                continue
                
            session = utils.parse_test_json(f'{args.user_path}/Word_test_output/{set_name}/{session_path}')
            turn_num = 0
            
            for turn in tqdm(session):
                turn_num += 1
                turn_id, instruction, label_api, reply, pred_api, pred_doc_path, label_doc_path, prompt_path = turn
                
                # 计算成本
                api_costs.append(len(pred_api))
                token_costs.append(calc_token_cost(f'{args.user_path}/{prompt_path}'))
                
                # 计算准确率
                str_c, pos_t, pos_c = calc_acc(
                    f'{args.user_path}/{label_doc_path}',
                    f'{args.user_path}/{pred_doc_path}',
                    instruction
                )
                
                string_total += 1
                string_correct += str_c
                position_total += pos_t
                position_correct += pos_c
                
            turn_nums.append(turn_num)
            print(f'Session {sess_id}: {turn_nums}')
    
    # 计算统计结果
    try:
        avg_api_costs = sum(api_costs) / len(api_costs)
        avg_token_costs = sum(token_costs) / len(token_costs)
        string_acc = string_correct / string_total
        position_acc = position_correct / position_total
    except ZeroDivisionError:
        print("No valid samples found")
        return
        
    # 输出结果
    print(f"Average API cost: {avg_api_costs:.2f}")
    print(f"Average token cost: {avg_token_costs:.2f}")
    print(f"String accuracy: {string_correct}/{string_total} = {string_acc:.2%}")
    print(f"Position accuracy: {position_correct}/{position_total} = {position_acc:.2%}")

def get_error_case(args):
    """获取错误案例
    
    Args:
        args: 参数配置对象
    """
    with open(f'{args.dataset}_{args.exp_name}_error_info.txt', 'w+', encoding='utf-8') as f:
        if args.tf:
            for idx in sorted(os.listdir(args.save_path)):
                if not os.path.isdir(os.path.join(args.save_path, str(idx))):
                    continue
                    
                for step in sorted(os.listdir(os.path.join(args.save_path, idx))):
                    if '.' in step or not os.path.isdir(os.path.join(args.save_path, str(idx), str(step))):
                        continue
                        
                    current_path = os.path.join(args.save_path, str(idx), str(step))
                    pred_path = os.path.join(current_path, f"{args.exp_name}_after_pred.docx")
                    label_path = os.path.join(current_path, "after_label.docx")
                    instruction_path = os.path.join(current_path, "instruction.txt")
                    api_pred_path = os.path.join(current_path, f"{args.exp_name}_api_pred.txt")
                    api_label_path = os.path.join(current_path, f"api_label.txt")

                    # 计算准确率
                    str_c, _, _ = calc_acc(label_path, pred_path, open(instruction_path, 'r').read().strip())
                    
                    # 如果预测错误,记录相关信息
                    if str_c != 1:
                        f.write(f"{idx}/{step}\n")
                        f.write(open(instruction_path, 'r').read().strip() + '\n')
                        f.write('## Label ##\n')
                        f.write(open(api_label_path, 'r').read().strip() + '\n')
                        f.write('## Prediction ##\n')
                        f.write(open(api_pred_path, 'r').read().strip() + '\n\n')
                        f.flush()

def calc_api_cost(api_pred_path, args):
    """计算API调用成本
    
    Args:
        api_pred_path: API预测结果文件路径
        args: 参数配置
        
    Returns:
        int: API调用数量
    """
    try:
        with open(api_pred_path, 'r') as f:
            api_calls = f.read().strip().split(';')
            return len([call for call in api_calls if call.strip()])
    except Exception as e:
        print(f"Error calculating API cost: {e}")
        return 0

def check_eval(args):
    """检查评估中的异常情况
    
    主要检查:
    1. API调用数量为0的案例
    2. 文件缺失的情况
    3. API调用失败的情况
    
    Args:
        args: 参数配置
    """
    token_costs, api_costs = [], []
    string_total = string_correct = position_total = position_correct = 0
    
    if args.tf:
        # 存储问题案例
        zero_api_cases = []  # API调用为0的案例
        missing_files = []   # 文件缺失的案例
        api_errors = []      # API调用错误的案例
        
        for idx in sorted(os.listdir(args.save_path)):
            if not os.path.isdir(os.path.join(args.save_path, str(idx))):
                continue
                
            for step in sorted(os.listdir(os.path.join(args.save_path, idx))):
                if '.' in step or not os.path.isdir(os.path.join(args.save_path, str(idx), str(step))):
                    continue
                    
                current_path = os.path.join(args.save_path, str(idx), str(step))
                
                # 检查所需文件
                files_to_check = {
                    'pred': f"{args.exp_name}_after_pred.docx",
                    'label': "after_label.docx",
                    'instruction': "instruction.txt",
                    'prompt': f"{args.exp_name}_prompt.txt",
                    'api_pred': f"{args.exp_name}_api_pred.txt"
                }
                
                missing = []
                for key, filename in files_to_check.items():
                    file_path = os.path.join(current_path, filename)
                    if not os.path.exists(file_path):
                        missing.append(filename)
                
                if missing:
                    missing_files.append({
                        'case': f'{idx}/{step}',
                        'missing': missing
                    })
                    continue
                
                # 检查API调用
                api_pred_path = os.path.join(current_path, f"{args.exp_name}_api_pred.txt")
                api_cost = calc_api_cost(api_pred_path, args)
                
                if api_cost == 0:
                    zero_api_cases.append(f'{idx}/{step}')
                
                try:
                    # 检查API调用结果
                    with open(api_pred_path, 'r') as f:
                        api_content = f.read().strip()
                    if 'error' in api_content.lower():
                        api_errors.append({
                            'case': f'{idx}/{step}',
                            'error': api_content
                        })
                except Exception as e:
                    api_errors.append({
                        'case': f'{idx}/{step}',
                        'error': str(e)
                    })
        
        # 输出检查结果
        print("\n=== Evaluation Check Results ===")
        
        if zero_api_cases:
            print("\nCases with zero API calls:")
            for case in zero_api_cases:
                print(f"- {case}")
                
        if missing_files:
            print("\nCases with missing files:")
            for case in missing_files:
                print(f"- {case['case']}: Missing {', '.join(case['missing'])}")
                
        if api_errors:
            print("\nCases with API errors:")
            for case in api_errors:
                print(f"- {case['case']}: {case['error']}")
                
        # 保存检查结果
        with open(f"{args.exp_name}_eval_check.txt", 'w') as f:
            f.write("=== Evaluation Check Results ===\n")
            
            f.write("\nCases with zero API calls:\n")
            for case in zero_api_cases:
                f.write(f"- {case}\n")
                
            f.write("\nCases with missing files:\n")
            for case in missing_files:
                f.write(f"- {case['case']}: Missing {', '.join(case['missing'])}\n")
                
            f.write("\nCases with API errors:\n")
            for case in api_errors:
                f.write(f"- {case['case']}: {case['error']}\n")



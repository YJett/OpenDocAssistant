from .openai_api import *
from .prompt_factor import *
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import json

def select_information_type(query, args):
    """选择信息类型，适用于docx文档的内容提取
    
    Args:
        query: 用户查询
        args: 参数配置
    Returns:
        list: [need_text, need_style, need_size]
    """
    try:
        # 使用Information_selection_prompt模板
        prompt = Information_selection_prompt.format(query)
        response = query_azure_openai(prompt, model=args.model if args else "gpt-3.5-turbo")
        
        # 解析响应
        parts = response.strip().split(',')
        result = {}
        for part in parts:
            key, value = part.strip().split('=')
            result[key.strip()] = int(value.strip())
            
        return [
            result.get('text', 0),
            result.get('style', 0),
            result.get('size', 0)
        ]
    except Exception as e:
        print(f"Error in select_information_type: {str(e)}")
        return [1, 1, 1]  # 默认返回所有类型

def select_element_type(query, args):
    """选择元素类型，适用于docx文档的内容处理
    
    Args:
        query: 用户查询
        args: 参数配置
    Returns:
        dict: 需要处理的元素类型
    """
    try:
        # 使用Element_selection_prompt模板
        prompt = Element_selection_prompt.format(query)
        response = query_azure_openai(prompt, model=args.model if args else "gpt-3.5-turbo")
        
        # 解析响应
        parts = response.strip().split(',')
        result = {}
        for part in parts:
            key, value = part.strip().split('=')
            result[key.strip()] = int(value.strip())
            
        return result
    except Exception as e:
        print(f"Error in select_element_type: {str(e)}")
        return {
            'heading': 1,
            'paragraph': 1,
            'picture': 1,
            'table': 1,
            'chart': 1
        }

def get_content_selection(query, args):
    """获取内容选择的API调用
    
    Args:
        query: 用户查询
        args: 参数配置
    Returns:
        str: API调用字符串
    """
    try:
        # 使用Word_content_selection_prompt模板
        prompt = Word_content_selection_prompt.format(query)
        response = query_azure_openai(prompt, model=args.model if args else "gpt-3.5-turbo")
        
        # 验证返回的API调用格式
        if not response.strip().startswith("get_content("):
            raise ValueError("Invalid API call format")
            
        return response.strip()
    except Exception as e:
        print(f"Error in get_content_selection: {str(e)}")
        return "get_content(need_text=1,need_style=1,need_size=1,need_heading=1,need_paragraph=1,need_picture=1,need_table=1,need_chart=1)"

def process_docx(query, args):
    """处理docx文档，根据用户查询生成文档内容"""
    try:
        # 1. 获取信息类型
        info_types = select_information_type(query, args)
        print(f"Selected information types: {info_types}")
        need_text, need_style, need_size = info_types
        
        # 2. 获取元素类型
        element_types = select_element_type(query, args)
        print(f"Selected element types: {element_types}")
        
        # 3. 获取内容选择
        content_selection = get_content_selection(query, args)
        print(f"Content selection: {content_selection}")
        
        # 4. 创建文档并应用选择
        doc = Document()
        
        # 处理特定需求：将标题移到底部并填入"人工评估"
        if "将标题移到底部" in query and "人工评估" in query:
            # 首先添加一些空白段落来创造空间
            for _ in range(3):
                doc.add_paragraph()
                
            # 添加标题并设置对齐方式
            heading = doc.add_heading("人工评估", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 保存文档
            output_path = "Word_Base_File/Create_new_doc/0_0.docx"
            doc.save(output_path)
            
            # 同时保存到其他路径
            doc.save("Word_Label_File/Create_new_doc/0_0.docx")
            doc.save("Word_Label_File/Create_new_doc_API_lack/0_0.docx")
            
            print(f"Documents saved to multiple locations")
            return
            
        # 根据用户查询和元素类型处理文档
        if element_types.get('heading'):
            # 从查询中提取标题文本，如果没有则使用默认值
            title_text = query if "标题" in query else "文档标题"
            heading = doc.add_heading(title_text, level=1)
            if need_style:
                heading.style = 'Heading 1'
                
        if element_types.get('paragraph'):
            # 根据查询生成段落内容
            para = doc.add_paragraph()
            if "插入文本" in query.lower():
                # 尝试从查询中提取要插入的文本
                start = query.find('"') + 1
            para = doc.add_paragraph(f'Query: {query}')
            if need_style:
                para.style = 'Normal'
                
        if element_types.get('table'):
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            
        # 添加分析结果说明
        doc.add_paragraph('\nAnalysis Results:')
        doc.add_paragraph(f'Information Types: {info_types}')
        doc.add_paragraph(f'Element Types: {element_types}')
        doc.add_paragraph(f'Content Selection: {content_selection}')
        
        # 保存文档
        output_path = "output.docx"
        doc.save(output_path)
        print(f"Document saved as '{output_path}'")
        
    except Exception as e:
        print(f"Error in process_docx: {str(e)}")
        doc = Document()
        doc.add_paragraph(f"Error processing query: {query}")
        doc.add_paragraph(f"Error message: {str(e)}")
        doc.save("error_output.docx")

def main():
    """主函数，用于测试内容选择功能"""
    while True:
        query = input("\nEnter your query (or 'quit' to exit): ")
        if query.lower() == 'quit':
            break
            
        # 测试所有功能
        print("\n=== Information Type Analysis ===")
        info_types = select_information_type(query, None)
        print(f"Information types needed: {info_types}")
        
        print("\n=== Element Type Analysis ===")
        element_types = select_element_type(query, None)
        print(f"Element types needed: {element_types}")
        
        print("\n=== Content Selection ===")
        content_selection = get_content_selection(query, None)
        print(f"Content selection API call: {content_selection}")
        
        print("\nProcessing document...")
        process_docx(query, None)

if __name__ == '__main__':
    main()

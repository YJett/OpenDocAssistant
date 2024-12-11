import collections
import collections.abc
from docx import Document
from docx.shared import RGBColor
import logging

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局变量
global_args = None
doc = None

# 形状类型列表
element_list = ['PARAGRAPH', 'TABLE', 'PICTURE', 'TEXT_BOX', 'AUTO_SHAPE']

# 全局变量
current_document = None

def set_document(doc):
    """设置当前操作的文档"""
    global current_document
    current_document = doc
    return current_document

def get_document():
    """获取当前操作的文档"""
    global current_document
    return current_document

def init_docx():
    """初始化一个新的Word文档"""
    global current_document
    current_document = Document()
    return current_document

def check_element_position(paragraph, element_type, position):
    """检查元素位置是否正确"""
    try:
        if element_type == 'text':
            return check_text_position(paragraph, position)
        elif element_type == 'image':
            return check_image_position(paragraph, position)
        else:
            logger.warning(f"Unknown element type: {element_type}")
            return False
    except Exception as e:
        logger.error(f"Error checking element position: {e}")
        return False

def get_document_content(doc):
    """获取文档内容"""
    try:
        content = []
        for para in doc.paragraphs:
            content.append(para.text)
        return '\n'.join(content)
    except Exception as e:
        logger.error(f"Error getting document content: {e}")
        return ""

# 获取段落内容
def get_paragraph_style(paragraph):
    style_info = []
    for run in paragraph.runs:
        font = run.font
        style_info.append({
            'bold': font.bold,
            'italic': font.italic,
            'underline': font.underline,
            'size': font.size,
            'color': font.color.rgb if font.color else None,
            'font_name': font.name,
            'alignment': paragraph.alignment
        })
    return style_info

# 获取图片的内容
# 获取图片信息
def get_picture_info(paragraph):
    pictures = []
    for run in paragraph.runs:
        # 检查 run._r 是否包含图形（图片）
        if hasattr(run._r, 'graphic') and run._r.graphic is not None:
            pic = run._r.graphic.graphicData.pic
            pictures.append({
                'width': pic.ext.cx,
                'height': pic.ext.cy
            })
    return pictures



# 获取文档内容
def get_content(need_text, need_style, need_position, need_picture, need_table):
    global doc
    content = ""

    # 遍历文档中的段落
    for paragraph in doc.paragraphs:  # 使用 doc.paragraphs 获取段落
        content += f"Paragraph: {paragraph.text}\n"

        if need_style:
            content += f"Style: {get_paragraph_style(paragraph)}\n"

        if need_picture:
            content += f"Pictures: {get_picture_info(paragraph)}\n"

    # 处理表格
    for table in doc.tables:
        content += f"Table:\n{get_table_info(table)}\n"

    return content

# 获取表格的内容
def get_table_info(table):
    table_info = ""
    for row in table.rows:
        row_data = "|"
        for cell in row.cells:
            row_data += f"{cell.text}|"
        row_data += "\n"
        table_info += row_data
    return table_info


# 获取段落中的所有文本
def get_paragraph_text(paragraph):
    return paragraph.text


# 获取文档内容



# 根据指令获取文档内容
def get_content_by_instructions(docx_path, instruction, args):
    global doc
    doc = Document(docx_path)

    content = f"There are {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables in this document.\n"

    # 获取内容
    ppt_content = get_content(
        need_text=1,
        need_style=1,
        need_position=1,
        need_picture=1,
        need_table=1
    )
    content += ppt_content
    return content


# 用于选择文档信息的示例函数
def select_information_type(instruction, args):
    need_text = 1
    need_style = 1
    need_position = 1
    return need_text, need_style, need_position


# 调用方式示例
if __name__ == "__main__":
    # docx_path = "/Users/ywootae/Desktop/DocxAPI/src/example.docx"
    # instruction = "Extract text and style from document"
    # args = None  # 可以传递其他参数（例如，从命令行或配置中获取）

    # result = get_content_by_instructions(docx_path, instruction, args)
    # print(result)
    pass

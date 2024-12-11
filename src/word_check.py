import sys
import os
import json
from src import word_reader
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING
from tqdm import tqdm

def check_paragraph_before(A, B):
    """检查段落A的行距是否小于段落B的行距
    Args:
        A: 段落A
        B: 段落B
    Returns:
        int: 如果A的行距小于B的行距返回1,否则返回0
    """
    return int(A.paragraph_format.line_spacing_rule < B.paragraph_format.line_spacing_rule)

def check_paragraph_after(A, B):
    """检查段落A的行距是否大于段落B的行距
    Args:
        A: 段落A
        B: 段落B
    Returns:
        int: 如果A的行距大于B的行距返回1,否则返回0
    """
    return int(A.paragraph_format.line_spacing_rule > B.paragraph_format.line_spacing_rule)

def check_heading_level(A, level):
    """检查段落A是否为指定级别的标题
    Args:
        A: 段落A
        level: 标题级别
    Returns:
        int: 如果A是指定级别的标题返回1,否则返回0
    """
    return int(A.style.name.startswith('Heading') and int(A.style.name[-1]) == level)

def choose_paragraph(doc, idx):
    """选择文档中指定索引的段落
    Args:
        doc: Word文档对象
        idx: 段落索引
    Returns:
        段落对象或None(如果索引无效)
    """
    paragraphs = doc.paragraphs
    if idx < len(paragraphs):
        return paragraphs[idx]
    return None

def choose_heading(doc, level):
    """选择文档中指定级别的第一个标题
    Args:
        doc: Word文档对象
        level: 标题级别
    Returns:
        标题段落对象或None(如果未找到)
    """
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and int(para.style.name[-1]) == level:
            return para
    return None

def choose_table(doc):
    """选择文档中的第一个表格
    Args:
        doc: Word文档对象
    Returns:
        表格对象或None(如果文档中没有表格)
    """
    if len(doc.tables) > 0:
        return doc.tables[0]
    return None

def choose_table_cell(table, row_id, col_id):
    """选择表格中指定位置的单元格
    Args:
        table: 表格对象
        row_id: 行索引
        col_id: 列索引
    Returns:
        单元格对象或None(如果索引无效)
    """
    if row_id < len(table.rows) and col_id < len(table.rows[0].cells):
        return table.cell(row_id, col_id)
    return None

def choose_picture(doc, idx=0):
    """选择文档中指定索引的图片
    Args:
        doc: Word文档对象
        idx: 图片索引,默认为0
    Returns:
        包含图片的run对象或None(如果未找到)
    """
    cur_idx = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if len(run._element.xpath('.//pic:pic')) > 0:
                if cur_idx == idx:
                    return run
                cur_idx += 1
    return None

def choose_object(doc, object_name):
    """根据对象名称选择文档中的对象
    Args:
        doc: Word文档对象
        object_name: 对象名称(如'paragraph0','heading1','table','picture0')
    Returns:
        对应的文档对象或None(如果未找到)
    """
    if 'paragraph' in object_name:
        para_id = int(object_name[10:]) if object_name != 'paragraph' else 0
        return choose_paragraph(doc, para_id)
    elif 'heading' in object_name:
        level = int(object_name[8:]) if object_name != 'heading' else 1
        return choose_heading(doc, level)
    elif object_name == 'table':
        return choose_table(doc)
    elif 'picture' in object_name:
        pic_id = int(object_name[8:]) if object_name != 'picture' else 0
        return choose_picture(doc, pic_id)
    else:
        return None

def check(doc, A, B, rel):
    """检查文档中两个对象之间的关系是否满足要求
    Args:
        doc: Word文档对象
        A: 第一个对象的名称
        B: 第二个对象的名称
        rel: 需要检查的关系
    Returns:
        int: 如果满足关系要求返回1,否则返回0
    """
    result = 1
    if B == "document":
        A = choose_object(doc, A)
        if A is None:
            print("No object found")
            return 0
        if "heading" in rel:
            level = int(rel[-1])
            result *= check_heading_level(A, level)
    else:
        A = choose_object(doc, A)
        B = choose_object(doc, B)
        if A is None or B is None:
            print("No object found") 
            return 0
        if "before" in rel:
            result *= check_paragraph_before(A, B)
        if "after" in rel:
            result *= check_paragraph_after(A, B)
    return result
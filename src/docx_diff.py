from docx import Document
from docx.enum.section import WD_HEADER_FOOTER
from docx.shared import RGBColor
import difflib

def compare_docx(doc1_path, doc2_path):
    """比较两个Word文档的主要差异"""
    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)
    differences = []

    # 比较基本结构
    compare_document_structure(doc1, doc2, differences)
    
    # 比较页眉页脚
    compare_headers_footers(doc1, doc2, differences)
    
    # 比较正文内容
    compare_body_content(doc1, doc2, differences)
    
    # 比较图片/图表
    compare_inline_shapes(doc1, doc2, differences)

    return differences

def compare_document_structure(doc1, doc2, differences):
    """比较文档基础结构"""
    # 比较段落总数
    if len(doc1.paragraphs) != len(doc2.paragraphs):
        differences.append(f"段落总数不一致：{len(doc1.paragraphs)} vs {len(doc2.paragraphs)}")

    # 比较表格总数
    if len(doc1.tables) != len(doc2.tables):
        differences.append(f"表格总数不一致：{len(doc1.tables)} vs {len(doc2.tables)}")

    # 比较章节数量
    if len(doc1.sections) != len(doc2.sections):
        differences.append(f"章节数量不一致：{len(doc1.sections)} vs {len(doc2.sections)}")

def compare_headers_footers(doc1, doc2, differences):
    """比较页眉页脚内容"""
    for section_idx, (s1, s2) in enumerate(zip(doc1.sections, doc2.sections)):
        # 比较页眉
        compare_header_footer(s1.header, s2.header, f"第{section_idx+1}节页眉", differences)
        # 比较页脚
        compare_header_footer(s1.footer, s2.footer, f"第{section_idx+1}节页脚", differences)

def compare_header_footer(hf1, hf2, hf_name, differences):
    """比较单个页眉/页脚"""
    if len(hf1.paragraphs) != len(hf2.paragraphs):
        differences.append(f"{hf_name}段落数量不一致：{len(hf1.paragraphs)} vs {len(hf2.paragraphs)}")
        return

    for p_idx, (p1, p2) in enumerate(zip(hf1.paragraphs, hf2.paragraphs)):
        compare_paragraph(p1, p2, f"{hf_name}第{p_idx+1}段", differences)

def compare_body_content(doc1, doc2, differences):
    """比较正文内容"""
    # 比较段落
    for p_idx, (p1, p2) in enumerate(zip(doc1.paragraphs, doc2.paragraphs)):
        compare_paragraph(p1, p2, f"第{p_idx+1}段", differences)
    
    # 比较表格
    for t_idx, (t1, t2) in enumerate(zip(doc1.tables, doc2.tables)):
        compare_table(t1, t2, f"第{t_idx+1}个表格", differences)

def compare_paragraph(p1, p2, para_name, differences):
    """比较单个段落"""
    # 文本内容比较
    if p1.text != p2.text:
        differences.append(f"{para_name}文本内容不同")
        # 使用difflib显示差异细节
        diff = difflib.ndiff(p1.text.splitlines(), p2.text.splitlines())
        differences.extend([f"{para_name}差异细节："] + list(diff))

    # 段落样式比较
    if p1.style.name != p2.style.name:
        differences.append(f"{para_name}样式不同：{p1.style.name} vs {p2.style.name}")

    # 段落格式比较
    compare_paragraph_format(p1.paragraph_format, p2.paragraph_format, para_name, differences)

    # 文字格式比较
    compare_runs(p1.runs, p2.runs, para_name, differences)

def compare_paragraph_format(pf1, pf2, para_name, differences):
    """比较段落格式"""
    formats = [
        ("对齐方式", pf1.alignment, pf2.alignment),
        ("左缩进", pf1.left_indent, pf2.left_indent),
        ("行距", pf1.line_spacing, pf2.line_spacing),
        ("段前间距", pf1.space_before, pf2.space_before),
        ("段后间距", pf1.space_after, pf2.space_after)
    ]
    
    for name, v1, v2 in formats:
        if v1 != v2:
            differences.append(f"{para_name}{name}不同：{v1} vs {v2}")

def compare_runs(runs1, runs2, para_name, differences):
    """比较文字格式"""
    # 在Word文档中，每次文本格式发生变化时（如字体/颜色/加粗等变化），就会创建一个新的Run(分段)
    if len(runs1) != len(runs2):
        differences.append(f"{para_name}文本分段数量不同：{len(runs1)} vs {len(runs2)}")
        return

    for r_idx, (r1, r2) in enumerate(zip(runs1, runs2)):
        # 字体名称
        if r1.font.name != r2.font.name:
            differences.append(f"{para_name}第{r_idx+1}段文字字体不同：{r1.font.name} vs {r2.font.name}")
        
        # 字体大小
        if r1.font.size != r2.font.size:
            differences.append(f"{para_name}第{r_idx+1}段文字大小不同：{r1.font.size} vs {r2.font.size}")

        # 字体颜色
        color1 = r1.font.color.rgb
        color2 = r2.font.color.rgb
        if color1 != color2:
            c1 = color1 if color1 else "默认"
            c2 = color2 if color2 else "默认"
            differences.append(f"{para_name}第{r_idx+1}段文字颜色不同：{c1} vs {c2}")

def compare_table(t1, t2, table_name, differences):
    """比较表格内容"""
    # 比较行列数
    if len(t1.rows) != len(t2.rows) or len(t1.columns) != len(t2.columns):
        differences.append(f"{table_name}结构不同：{len(t1.rows)}x{len(t1.columns)} vs {len(t2.rows)}x{len(t2.columns)}")
        return

    # 比较单元格内容
    for row_idx, (row1, row2) in enumerate(zip(t1.rows, t2.rows)):
        for col_idx, (cell1, cell2) in enumerate(zip(row1.cells, row2.cells)):
            prefix = f"{table_name}第{row_idx+1}行第{col_idx+1}列"
            # 比较文本内容
            if cell1.text != cell2.text:
                differences.append(f"{prefix}文本内容不同")
            
            # 比较单元格格式
            for p1, p2 in zip(cell1.paragraphs, cell2.paragraphs):
                compare_paragraph(p1, p2, f"{prefix}段落", differences)

def compare_inline_shapes(doc1, doc2, differences):
    """比较图片/图表"""
    shapes1 = doc1.inline_shapes
    shapes2 = doc2.inline_shapes
    
    if len(shapes1) != len(shapes2):
        differences.append(f"图形数量不同：{len(shapes1)} vs {len(shapes2)}")
    else:
        for idx, (s1, s2) in enumerate(zip(shapes1, shapes2)):
            if s1.type != s2.type:
                differences.append(f"第{idx+1}个图形类型不同：{s1.type} vs {s2.type}")

if __name__ == "__main__":
    import argparse
    
    # 创建解析器
    parser = argparse.ArgumentParser(description="")

    # 添加两个字符串类型的参数
    parser.add_argument('--doc1', type=str, help="the first document")
    parser.add_argument('--doc2', type=str, help="the second document")

    # 解析命令行参数
    args = parser.parse_args()
    
    file1 = args.doc1
    file2 = args.doc2
    
    diffs = compare_docx(file1, file2)
    
    if diffs:
        print("发现差异：")
        for i, diff in enumerate(diffs, 1):
            print(f"{i}. {diff}")
    else:
        print("两个文档内容完全一致")